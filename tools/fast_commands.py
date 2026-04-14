import re
import os
import random
import unicodedata
from . import pc_control
from .code_executor import execute_code


def _normalize(text: str) -> str:
    """Lowercase + quita acentos + colapsa espacios. Sirve para que todos los
    patterns funcionen sin acentos y cubran mas variaciones de STT.
    Ejemplo: 'Búscame algo en YouTube' -> 'buscame algo en youtube'
    """
    if not text:
        return ""
    t = text.lower().strip()
    # NFD: descompone caracteres acentuados; luego filtra diacriticos
    t = unicodedata.normalize("NFD", t)
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    # Colapsar espacios multiples
    t = re.sub(r"\s+", " ", t)
    return t


class FastCommandDetector:
    """
    Detecta comandos simples y los ejecuta INMEDIATAMENTE sin pasar por la IA.
    Solo los comandos que no necesitan razonamiento pasan por aqui.
    """

    # Respuestas variadas, con personalidad tipo JARVIS — tono seco, britanico,
    # con sarcasmo ocasional sutil. Mezcla de confirmaciones formales y
    # comentarios laterales para que no suene robotico ni monocorde.
    RESPONSES = {
        "open": [
            "Enseguida, senor.",
            "Por supuesto.",
            "A su disposicion.",
            "Como desee.",
            "Procedo.",
            "Inmediatamente, senor.",
            "Faltaba mas.",
            "En ello.",
            "Abriendo. Intentare contener mi entusiasmo.",
            "Hecho. Si me permite el comentario, una eleccion predecible.",
            "Como cada dia, senor.",
            "A la orden.",
        ],
        "close": [
            "Cerrando, senor.",
            "Hecho.",
            "Considereselo cerrado.",
            "Por fin, diria yo.",
            "Listo.",
            "Fuera de escena.",
            "Cerrado. Su estabilidad mental, presumo, lo agradece.",
            "A su voluntad, senor.",
            "Asunto zanjado.",
        ],
        "media": [
            "Como desee, senor.",
            "Procedo.",
            "A la orden.",
            "Hecho.",
            "Enseguida.",
            "Sus gustos, senor. No los juzgo.",
            "Reproduciendo.",
        ],
        "search": [
            "Buscando, senor.",
            "Indagando.",
            "A ver que ha hecho la humanidad hoy al respecto.",
            "Consultando.",
            "Permitame.",
            "Enseguida.",
            "Investigando. Esperemos algo util.",
        ],
        "volume": [
            "Volumen ajustado, senor.",
            "Hecho.",
            "Como prefiera.",
            "Listo.",
            "A su oido, senor.",
        ],
        "system": [
            "Hecho, senor.",
            "Ejecutado.",
            "A su disposicion.",
            "Procedo.",
            "Considereselo hecho.",
            "Como ordene.",
        ],
    }

    # Verbos sinonimos para abrir aplicaciones (regex-grupo, no-capturador)
    # Cubre: abre, abreme, ejecuta, ejecutame, lanza, lanzame, inicia, iniciame,
    # arranca, corre, prende, enciende, activa, levanta, dame, sacame, ponme,
    # mas formas en infinitivo (abrir, ejecutar, etc.)
    # Nota: 'ponme' se omite aqui porque es ambiguo ("ponme una cancion" no es
    # abrir). El compound pattern ya lo captura cuando viene con "y pon/busca".
    _OPEN_VERBS = (
        r"(?:abre|abreme|abrir|abrime|abrila|abrilo|ejecuta|ejecutame|ejecutar|"
        r"lanza|lanzame|lanzar|lanzalo|lanzala|inicia|iniciame|iniciar|iniciame|"
        r"arranca|arrancame|arrancar|corre|correme|correr|"
        r"prende|prendeme|prender|enciende|enciendeme|encender|activa|activame|activar|"
        r"levanta|levantame|levantar|dame|sacame|abremela|abremelo|"
        r"muestrame|mostrame|pasame|pone\s+a|pon\s+a|"
        r"podrias\s+abrir|puedes\s+abrir|puedes\s+iniciar|necesito|quiero|quisiera|"
        r"abre\s+por\s+favor|por\s+favor\s+abre|open)"
    )
    # Verbos sinonimos para cerrar
    _CLOSE_VERBS = (
        r"(?:cierra|cierrame|cierralo|cierramela|cerrar|cerrame|cerralo|"
        r"termina|terminame|terminar|finaliza|finalizame|finalizar|"
        r"mata|matame|matar|kill|quita|quitame|quitar|apaga(?!\s+(?:el|la|mi|pc))|"
        r"close|saca|sacame)"
    )
    # Verbos sinonimos para reproducir (compuestos con spotify/youtube)
    _PLAY_VERBS = (
        r"(?:pon|ponme|ponle|ponelo|ponela|poneme|reproduce|reproduceme|reproducir|"
        r"reprodu|coloca|colocame|colocalo|busca|buscame|buscalo|buscala|play|dale\s+play|"
        r"escucha|escuchame|escuchalo|tocame|toca|tocala|tocalo|tira|tirame|"
        r"sueltame|suelta|dame|metele|meteme|cantame|dale|"
        r"ve|mira|miralo|mostrame|muestrame|abreme|abre|quiero\s+ver|"
        r"quiero\s+escuchar|quiero\s+oir|hazme\s+sonar)"
    )

    # Patrones de comandos rapidos
    # FAST COMMANDS: patrones que cubren acciones simples sin razonamiento.
    # Ordenados de MAS especifico a MAS general — el primero que matchea gana.
    PATTERNS = [
        # === ACCIONES COMPUESTAS GENERICAS (deben ir ANTES del open_app simple) ===
        # "abre/ejecuta/lanza <APP> y <VERBO> <CONTENIDO>"
        # Captura: (app, verbo_accion, contenido). El handler decide que tool usar.
        # Ejemplos:
        #   "ejecuta spotify y pon EL REGENTE"  -> spotify_search(EL REGENTE)
        #   "abre youtube y pon Bad Bunny"      -> open_and_search(youtube, Bad Bunny)
        #   "lanza chrome y busca recetas"      -> open_and_search(chrome, recetas) o search_web
        (rf"^{_OPEN_VERBS}\s+(?:la\s+|el\s+|los\s+|las\s+|mi\s+|me\s+)?(?:app\s+|aplicacion\s+|programa\s+|web\s+|pagina\s+)?(.+?)(?:\s+y\s+|\s*,\s*)({_PLAY_VERBS}|busca|buscame|search|escribe|escribeme)\s+(.+)", "_handle_open_compound"),

        # === SPOTIFY EXPLICITO (solo cuando menciona "spotify" textualmente) ===
        (r"(?:pon|reproduce|reproducir|play|busca|coloca|ponme).+(?:me gusta|liked|favorit|mis gustos).*(?:en\s+)?spotify", "_handle_spotify_liked"),
        (r"(?:busca|buscar|buscame)\s+en\s+spotify\s+(.+)", "_handle_spotify_search"),
        (r"(?:pon|reproduce|play|coloca|ponme)\s+(.+?)\s+en\s+spotify", "_handle_spotify_search"),
        (r"en\s+spotify\s+(?:pon|reproduce|busca|play)\s+(.+)", "_handle_spotify_search"),
        (r"spotify\s+(?:pon|reproduce|busca|play)\s+(.+)", "_handle_spotify_search"),

        # === BUSQUEDA/REPRODUCCION EN SITIOS WEB ESPECIFICOS ===
        # "busca/reproduce/pon en <sitio> X"
        (r"^(?:busca|buscar|buscame|buscalo|buscala|search|googlea|reproduce|reproduceme|reproducir|pon|ponme|ponle|coloca|colocame|play|muestrame|mostrame)\s+en\s+(youtube|google|github|twitter|x|reddit|netflix|amazon|mercadolibre|twitch|linkedin)\s+(.+)", "_handle_search_on_site"),
        # "en youtube busca/pon/reproduce X", "youtube busca X"
        (r"^(?:en\s+)?(youtube|github|twitter|reddit|netflix|amazon|twitch|linkedin|mercadolibre)\s+(?:busca|buscame|pon|ponme|reproduce|coloca|search|muestrame|mostrame)\s+(.+)", "_handle_search_on_site"),
        # "busca/reproduce/pon X en <sitio>" (query antes del sitio)
        (r"^(?:busca|buscar|buscame|buscalo|reproduce|reproduceme|pon|ponme|coloca|colocame|play)\s+(.+?)\s+en\s+(youtube|google|github|twitter|x|reddit|netflix|amazon|mercadolibre|twitch|linkedin)$", "_handle_search_on_site_reverse"),

        # === BUSQUEDA WEB GENERICA (Google) ===
        # "busca X", "busca en google X", "googlea X", "busca en internet X"
        (r"^(?:busca|buscame|buscalo|buscala|search|googlea)\s+(?:en\s+(?:google|internet|la\s+web|la\s+red)\s+)?(.+)", "_handle_search"),
        # Media controls (inequivocos, no necesitan contexto)
        (r"(?:pon|reproduce|play|reproducir)\s+(?:la\s+)?musica$", "_handle_play"),
        (r"(?:pausa|pausar|pause|para|detener|deten|stop|detente)(?:\s+(?:la\s+)?(?:musica|cancion|reproduccion|todo))?$", "_handle_play"),
        (r"(?:siguiente|next|salta|skip)(?:\s+(?:cancion|track|tema))?$", "_handle_next"),
        (r"(?:anterior|previous|prev|atras)(?:\s+(?:cancion|track|tema))?$", "_handle_prev"),

        # === REPRODUCIR X (sin plataforma explicita) — DECIDIR spotify/youtube ===
        # "reproduce el regente", "pon bad bunny", "cantame una cancion de...", etc.
        # Debe ir DESPUES de patrones spotify/youtube explicitos (ya capturados arriba)
        # y DESPUES de "reproduce musica" (media play/pause).
        (r"^(?:reproduce|reproduceme|reproducir|pon|ponme|ponle|coloca|colocame|play|cantame|escuchame|tocame)\s+(.+)", "_handle_play_smart"),
        # Volumen (inequivoco)
        (r"(?:sube|subir|subele|aumenta).+volumen.*?(\d+)?", "_handle_volume_up"),
        (r"(?:baja|bajar|bajale|reduce|disminuye).+volumen.*?(\d+)?", "_handle_volume_down"),
        (r"(?:volumen|volume)\s+(?:a|al|en)?\s*(\d+)", "_handle_volume_set"),
        (r"(?:silencia|mute|mutea|silenciar)$", "_handle_mute"),
        # Hora (inequivoco)
        (r"(?:que\s+hora|hora\s+es|dime\s+la\s+hora|que\s+hora\s+es)", "_handle_time"),
        # Captura de pantalla (inequivoco)
        (r"(?:captura|screenshot|pantallazo|captura\s+de\s+pantalla)", "_handle_screenshot"),
        # Grabar pantalla
        (r"(?:graba|grabar|record)\s+(?:los\s+)?(?:ultimos|últimos)\s+(\d+)\s+(?:segundos|seg|s)", "_handle_record"),
        (r"(?:graba|grabar|record)\s+(?:la\s+)?(?:pantalla|screen)(?:\s+(\d+)\s+(?:segundos|seg|s))?", "_handle_record"),
        (r"(?:graba|grabar|record)\s+(\d+)\s+(?:segundos|seg|s)", "_handle_record"),
        # Bloquear PC (inequivoco)
        (r"(?:bloquea|bloquear|lock)\s+(?:el\s+|la\s+|mi\s+)?(?:pc|computador|equipo|computadora|ordenador|compu|laptop|portatil|pantalla|sesion|screen)", "_handle_lock"),
        (r"bloquea(?:lo|la|te)?$", "_handle_lock"),
        # Apagar / reiniciar (inequivoco)
        (r"(?:apaga|apagar)\s+(?:el\s+)?(?:pc|computador|equipo|computadora|ordenador)", "_handle_shutdown"),
        (r"(?:apaga|apagar)(?:lo|te)?$", "_handle_shutdown"),
        (r"(?:reinicia|reiniciar|restart)\s+(?:el\s+)?(?:pc|computador|equipo|computadora|ordenador)", "_handle_restart"),
        (r"(?:reinicia|reiniciar|restart)(?:lo|te)?$", "_handle_restart"),
        (r"(?:suspende|suspender|dormir|duerme|sleep)\s+(?:el\s+)?(?:pc|computador|equipo|computadora|ordenador)", "_handle_sleep"),
        (r"(?:hiberna|hibernar|hibernate)\s+(?:el\s+)?(?:pc|computador|equipo|computadora|ordenador)", "_handle_hibernate"),
        (r"(?:cierra|cerrar)\s+(?:la\s+)?sesion", "_handle_logoff"),
        (r"(?:cancela|cancelar)\s+(?:el\s+)?(?:apagado|reinicio|shutdown|restart)", "_handle_cancel_shutdown"),
        # Presentacion (inequivoco)
        (r"(?:presentate|preséntate|quien\s+eres|quién\s+eres|como\s+te\s+llamas|cómo\s+te\s+llamas|que\s+eres|qué\s+eres|dime\s+quien\s+eres|tu\s+nombre)", "_handle_introduce"),

        # === ABRIR / CERRAR APP (catch-all al FINAL) ===
        # Abrir cualquier app: "abre X", "ejecuta X", "lanza X", "inicia X", etc.
        # Se evalua despues de los patrones especificos (spotify, youtube, etc.)
        # Si la app no existe, _handle_open retorna None y cae al IA.
        (rf"^{_OPEN_VERBS}\s+(?:la\s+|el\s+|los\s+|las\s+|mi\s+|me\s+)?(?:app\s+|aplicacion\s+|programa\s+)?(.+)", "_handle_open"),
        # Cerrar app (excluye pc/sesion que tienen sus propios patrones de shutdown/logoff arriba)
        (rf"^{_CLOSE_VERBS}\s+(?:la\s+|el\s+|los\s+|las\s+|mi\s+|me\s+)?(?:app\s+|aplicacion\s+|programa\s+|ventana\s+(?:de\s+)?)?(?!(?:el\s+|la\s+)?(?:pc|computador|computadora|equipo|ordenador|sesion|sesión)\b)(.+)", "_handle_close"),
    ]

    def __init__(self):
        self._compiled = [(re.compile(p, re.IGNORECASE), h) for p, h in self.PATTERNS]

    def _random_response(self, category: str) -> str:
        return random.choice(self.RESPONSES.get(category, self.RESPONSES["system"]))

    def try_execute(self, user_input: str) -> tuple[bool, str]:
        """
        Intenta ejecutar un comando rapido.
        Retorna (handled: bool, response: str).
        Si handled es False, el mensaje debe ir a la IA.
        """
        # Normalizar: lowercase + sin acentos + espacios colapsados. Esto
        # garantiza que los patterns (todos sin acentos) cubran cualquier
        # variante de STT o teclado ("Búscame", "BUSCAME", "buscame").
        text = _normalize(user_input)

        # Quitar "jarvis" del inicio (el usuario puede decir "Jarvis abre spotify")
        text_clean = re.sub(r'^(?:jarvis|hey\s+jarvis|oye\s+jarvis|oye|hey)[,.\s]*', '', text).strip()
        if text_clean:
            text = text_clean

        for pattern, handler_name in self._compiled:
            match = pattern.search(text)
            if match:
                handler = getattr(self, handler_name)
                result = handler(match)
                if result is not None:
                    return True, result

        return False, ""

    # Mapa de sitios web conocidos para resolver app -> URL en handlers compuestos
    _WEB_KEYWORDS = {
        "youtube": "youtube.com",
        "google": "google.com",
        "gmail": "mail.google.com",
        "twitter": "twitter.com",
        "x": "x.com",
        "facebook": "facebook.com",
        "instagram": "instagram.com",
        "github": "github.com",
        "whatsapp web": "web.whatsapp.com",
        "chatgpt": "chat.openai.com",
        "linkedin": "linkedin.com",
        "reddit": "reddit.com",
        "twitch": "twitch.tv",
        "netflix": "netflix.com",
        "amazon": "amazon.com",
        "mercadolibre": "mercadolibre.com",
    }

    def _handle_open(self, match) -> str | None:
        raw_target = match.group(1).strip().rstrip(".")
        # Limpia articulos ('la calculadora' -> 'calculadora', 'a chrome' -> 'chrome')
        target = pc_control._clean_app_name(raw_target) or raw_target
        target_lower = target.lower()

        for key, url in self._WEB_KEYWORDS.items():
            if key in target_lower:
                result = pc_control.open_website(url)
                if result["success"]:
                    return f"{self._random_response('open')} {key.capitalize()} abierto."
                return f"No pude abrir {key}."

        # Es una app - buscar en todo el sistema
        result = pc_control.find_and_open_app(target)
        if result["success"]:
            return f"{self._random_response('open')} {target.capitalize()} abierto."
        return None  # Dejar que la IA maneje si no se pudo

    # URLs de busqueda directa para sitios web conocidos — fallback garantizado
    # que nunca depende de automatizar teclado.
    _WEB_SEARCH_URLS = {
        "youtube": "https://www.youtube.com/results?search_query={}",
        "google": "https://www.google.com/search?q={}",
        "github": "https://github.com/search?q={}",
        "amazon": "https://www.amazon.com/s?k={}",
        "mercadolibre": "https://listado.mercadolibre.com/{}",
        "twitter": "https://twitter.com/search?q={}",
        "x": "https://x.com/search?q={}",
        "reddit": "https://www.reddit.com/search/?q={}",
        "linkedin": "https://www.linkedin.com/search/results/all/?keywords={}",
        "netflix": "https://www.netflix.com/search?q={}",
        "twitch": "https://www.twitch.tv/search?term={}",
    }

    def _handle_open_compound(self, match) -> str | None:
        """
        Maneja "abre/ejecuta <APP> y <VERBO> <CONTENIDO>" para CUALQUIER app.

        Nunca devuelve None cuando la app es reconocible: si la automatizacion
        de teclado falla, caemos a una URL de busqueda directa (mas rapido y
        100% fiable). Solo devuelve None para apps totalmente desconocidas
        donde no se puede garantizar nada util sin IA.
        """
        from . import automation
        import urllib.parse
        import time

        raw_app = match.group(1).strip().rstrip(",.")
        verb = match.group(2).strip().lower()
        content = match.group(3).strip().rstrip(".")
        if not content:
            return None

        # Limpia articulos
        app = pc_control._clean_app_name(raw_app) or raw_app
        app_lower = app.lower()

        # Verbo "escribe X" -> abrir app y tipear
        if verb in ("escribe", "escribeme"):
            open_result = pc_control.find_and_open_app(app)
            if not open_result.get("success"):
                # Probar como web
                for key, url in self._WEB_KEYWORDS.items():
                    if key in app_lower:
                        pc_control.open_website(url)
                        open_result = {"success": True}
                        break
            if not open_result.get("success"):
                return None
            time.sleep(1.5)  # esperar a que la app abra
            try:
                automation.type_text(content)
            except Exception:
                return None
            return f"{self._random_response('open')} {app.capitalize()} abierto y texto escrito."

        # Spotify: spotify_search ya abre Spotify automaticamente
        if "spotify" in app_lower:
            return self._spotify_search_and_play(content)

        # Sitios web conocidos: URL directa de busqueda (100% fiable, <1s)
        for key in self._WEB_SEARCH_URLS:
            if key in app_lower:
                url = self._WEB_SEARCH_URLS[key].format(urllib.parse.quote(content))
                result = pc_control.open_website(url)
                if result.get("success"):
                    return f"{self._random_response('open')} {key.capitalize()} abierto, buscando '{content}'."
                return f"No pude abrir {key} para buscar '{content}'."

        # App desktop (chrome, brave, firefox, edge) + busca/ve -> abrir app y
        # navegar a Google con la query.
        browser_apps = ("chrome", "brave", "firefox", "edge", "opera", "vivaldi", "tor")
        is_browser = any(b in app_lower for b in browser_apps)
        if is_browser or verb in ("busca", "buscame", "buscalo", "buscala", "search", "ve", "mira", "miralo"):
            # Intentar abrir la app; si falla, abrir la busqueda en navegador por defecto
            open_result = pc_control.find_and_open_app(app) if is_browser else {"success": False}
            if open_result.get("success"):
                time.sleep(0.8)
            pc_control.search_web(content)
            app_label = app.capitalize() if is_browser else "navegador"
            return f"{self._random_response('search')} {app_label} abierto, buscando '{content}'."

        # Si el verbo es play-like y la "app" parece un servicio de musica/video
        # desconocido, caer a YouTube como fallback razonable.
        if verb in ("pon", "ponme", "reproduce", "reproduceme", "reproducir", "play", "dale", "cantame"):
            url = self._WEB_SEARCH_URLS["youtube"].format(urllib.parse.quote(f"{content} {app}"))
            pc_control.open_website(url)
            return f"{self._random_response('media')} Buscando '{content}' en YouTube."

        # App desconocida + verbo ambiguo -> dejar al IA (que clasifique y decida)
        return None

    def _handle_whatsapp_send(self, match) -> str | None:
        contact = match.group(1).strip().rstrip(".,")
        message = match.group(2).strip().rstrip(".")
        # Quitar comillas accidentales
        contact = contact.strip('"\'')
        message = message.strip('"\'')
        if not contact or not message:
            return None
        result = pc_control.whatsapp_send_message(contact, message)
        if result["success"]:
            return f"Listo, senor. Mensaje enviado a {contact} por WhatsApp."
        return f"No pude enviar el mensaje: {result['message']}"

    def _handle_close(self, match) -> str:
        target = match.group(1).strip().rstrip(".")
        result = pc_control.close_application(target)
        if result["success"]:
            return f"{self._random_response('close')} {target.capitalize()} cerrado."
        return f"No encontre {target} ejecutandose."

    def _handle_search(self, match) -> str:
        query = match.group(1).strip().rstrip(".")
        # Limpiar prefijos comunes
        for prefix in ["en google ", "en internet ", "en la web "]:
            if query.startswith(prefix):
                query = query[len(prefix):]
        # Si queda vacio (el usuario solo dijo "busca en internet"), fallback
        if not query:
            return "¿Que desea que busque, senor?"
        pc_control.search_web(query)
        return f"{self._random_response('search')} Buscando: {query}."

    def _handle_search_on_site(self, match) -> str:
        """Busca en un sitio web especifico usando URL directa de busqueda.
        Grupo 1: sitio (youtube, google, etc.) — Grupo 2: query.
        """
        import urllib.parse
        site = match.group(1).strip().lower()
        query = match.group(2).strip().rstrip(".,;:!?")
        if not query:
            return f"¿Que desea buscar en {site.capitalize()}, senor?"
        url_template = self._WEB_SEARCH_URLS.get(site)
        if not url_template:
            # Fallback a Google
            pc_control.search_web(f"{query} site:{site}.com")
            return f"{self._random_response('search')} Buscando '{query}' en {site}."
        url = url_template.format(urllib.parse.quote(query))
        pc_control.open_website(url)
        return f"{self._random_response('search')} Buscando '{query}' en {site.capitalize()}."

    def _handle_search_on_site_reverse(self, match) -> str:
        """'busca X en youtube' — query primero, sitio segundo."""
        import urllib.parse
        query = match.group(1).strip().rstrip(".,;:!?")
        site = match.group(2).strip().lower()
        if not query:
            return f"¿Que desea buscar en {site.capitalize()}, senor?"
        url_template = self._WEB_SEARCH_URLS.get(site)
        if not url_template:
            pc_control.search_web(f"{query} site:{site}.com")
            return f"{self._random_response('search')} Buscando '{query}' en {site}."
        url = url_template.format(urllib.parse.quote(query))
        pc_control.open_website(url)
        return f"{self._random_response('search')} Buscando '{query}' en {site.capitalize()}."

    def _handle_play(self, match) -> str:
        pc_control.media_play_pause()
        return self._random_response("media")

    def _handle_next(self, match) -> str:
        pc_control.media_next()
        return f"{self._random_response('media')} Siguiente cancion."

    def _handle_prev(self, match) -> str:
        pc_control.media_previous()
        return f"{self._random_response('media')} Cancion anterior."

    def _handle_spotify_liked(self, match) -> str:
        # Buscar "liked songs" en Spotify usando el mismo metodo que funciona
        return self._spotify_search_and_play("liked songs")

    def _clean_spotify_query(self, raw: str) -> str:
        """Limpia el texto para obtener solo el nombre de la cancion/album/artista.

        Ejemplos:
        - "la cancion llamada Lloraras de Oscar d'leon" -> "Lloraras de Oscar d'leon"
        - "el album Thriller de Michael Jackson" -> "Thriller de Michael Jackson"
        - "musica de Bad Bunny" -> "Bad Bunny"
        - "algo de reggaeton" -> "reggaeton"
        - "Lloraras de Oscar d'leon en spotify" -> "Lloraras de Oscar d'leon"
        """
        query = raw.strip().strip('"\'.,;:!?')
        query_lower = query.lower()

        # Limpiar "en spotify" al final PRIMERO
        for suffix in [" en spotify", " de spotify", " en el spotify"]:
            if query_lower.endswith(suffix):
                query = query[:len(query) - len(suffix)].strip()
                query_lower = query.lower()

        # Prefijos que son puro ruido (el contenido real viene despues)
        noise_prefixes = [
            "la cancion llamada ", "la cancion que se llama ",
            "cancion llamada ", "cancion que se llama ",
            "el album llamado ", "el album que se llama ",
            "album llamado ", "album que se llama ",
            "el disco llamado ", "el disco que se llama ",
            "que se llama ", "llamada ", "llamado ",
        ]
        for prefix in noise_prefixes:
            if query_lower.startswith(prefix):
                query = query[len(prefix):]
                return query.strip().strip('"\'.,;:!?')

        # Prefijos que indican "busca esto DE artista" -> mantener todo
        # "la cancion X de Y" -> "X de Y" (mantener "de Y" porque es el artista)
        content_prefixes = [
            "la cancion ", "cancion ", "el tema ", "tema ",
            "el track ", "track ", "la song ", "song ",
            "el album ", "album ", "el disco ", "disco ",
            "la playlist ", "playlist ",
        ]
        for prefix in content_prefixes:
            if query_lower.startswith(prefix):
                query = query[len(prefix):]
                return query.strip().strip('"\'.,;:!?')

        # "musica de X" o "algo de X" -> solo X (el artista)
        for prefix in ["musica de ", "algo de ", "music de ", "canciones de "]:
            if query_lower.startswith(prefix):
                query = query[len(prefix):]
                return query.strip().strip('"\'.,;:!?')

        return query.strip().strip('"\'.,;:!?')

    def _spotify_search_and_play(self, query: str) -> str:
        """Busca en Spotify y lo reproduce."""
        if not query:
            return "No entendi que buscar en Spotify. Dime el nombre de la cancion o artista."

        result = pc_control.spotify_search_and_play(query)
        if result["success"]:
            return f"{self._random_response('media')} Reproduciendo '{query}' en Spotify."
        return f"No pude buscar en Spotify: {result['message']}"

    def _handle_spotify_search(self, match) -> str:
        query = self._clean_spotify_query(match.group(1))
        return self._spotify_search_and_play(query)

    def _handle_spotify_play(self, match) -> str:
        pc_control.media_play_pause()
        return f"{self._random_response('media')} Reproduciendo en Spotify."

    def _handle_spotify_catch_all(self, match) -> str:
        """Catch-all: cualquier 'reproduce/pon X' que no matcheo nada mas -> Spotify."""
        raw = match.group(1).strip()
        # Ignorar si es "musica" sola (eso es play/pause)
        if raw.lower().strip() in ["musica", "la musica", "music"]:
            pc_control.media_play_pause()
            return self._random_response("media")
        query = self._clean_spotify_query(raw)
        if query:
            return self._spotify_search_and_play(query)
        return None

    def _handle_play_smart(self, match) -> str | None:
        """Reproduce X sin plataforma — decide Spotify vs YouTube segun estado.

        Logica:
          - Si Spotify esta corriendo -> buscar y reproducir ahi.
          - Si no -> abrir URL directa de YouTube con la busqueda.
        Asi el usuario no tiene que especificar la plataforma cada vez.
        """
        import urllib.parse
        raw = match.group(1).strip().rstrip(".,;:!?")
        # "reproduce musica" ya lo cubre _handle_play (con anchor $). Si llega
        # aqui sin mas texto, tratarlo como play/pause.
        if not raw or raw.lower() in ("musica", "la musica", "music"):
            pc_control.media_play_pause()
            return self._random_response("media")

        # Limpiar prefijos ruido para no buscar "la cancion llamada X"
        query = self._clean_spotify_query(raw)
        if not query:
            return None

        if pc_control._is_spotify_running():
            return self._spotify_search_and_play(query)

        # Fallback: YouTube (URL directa, 100% fiable)
        url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(query)}"
        pc_control.open_website(url)
        return f"{self._random_response('media')} Spotify no esta abierto; buscando '{query}' en YouTube."

    def _handle_volume_up(self, match) -> str:
        level = match.group(1)
        if level:
            pc_control.set_volume(int(level))
            return f"{self._random_response('volume')} Volumen al {level}%."
        pc_control.set_volume(80)
        return f"{self._random_response('volume')} Volumen subido."

    def _handle_volume_down(self, match) -> str:
        level = match.group(1)
        if level:
            pc_control.set_volume(int(level))
            return f"{self._random_response('volume')} Volumen al {level}%."
        pc_control.set_volume(30)
        return f"{self._random_response('volume')} Volumen bajado."

    def _handle_volume_set(self, match) -> str:
        level = int(match.group(1))
        pc_control.set_volume(level)
        return f"{self._random_response('volume')} Volumen al {level}%."

    def _handle_mute(self, match) -> str:
        pc_control.mute_volume()
        return f"{self._random_response('media')} Sonido silenciado."

    def _handle_time(self, match) -> str:
        result = pc_control.get_datetime()
        return f"Son las {result['message']}, senor."

    def _handle_screenshot(self, match) -> str:
        result = pc_control.take_screenshot()
        if result["success"]:
            return f"{self._random_response('system')} {result['message']}"
        return "No pude tomar la captura."

    def _handle_record(self, match) -> str:
        seconds = 30
        for g in match.groups():
            if g and g.isdigit():
                seconds = int(g)
                break
        result = pc_control.record_screen(seconds)
        if result["success"]:
            return f"{self._random_response('system')} {result['message']}"
        return result["message"]

    def _handle_introduce(self, match) -> str:
        from config import Config
        name = Config.ASSISTANT_NAME
        intros = [
            f"{name}, senor. Su mayordomo digital. Una inteligencia artificial a su entera disposicion.",
            f"Soy {name}. Discretamente diseñado para facilitarle la vida. Pidame lo que sea razonable, y algunas cosas que no lo sean.",
            f"Me llamo {name}, senor. Inspirado en cierto mayordomo digital de Malibu, aunque con mejor sentido del humor, si me permite.",
            f"{name} a su servicio. Abro aplicaciones, consulto informacion, controlo su equipo. Y escucho, cuando usted lo necesita.",
            f"Soy {name}. Piense en mi como su asistente personal, con la diferencia de que yo no duermo ni pido vacaciones, senor.",
            f"{name}, senor. Mi proposito: ejecutar sus ordenes con discrecion y eficiencia. El resto es musica de fondo.",
        ]
        return random.choice(intros)

    def _handle_lock(self, match) -> str:
        pc_control.lock_pc()
        options = [
            "Bloqueando, senor. Hasta su regreso.",
            "Equipo asegurado.",
            "Cerrando puertas. Que tenga buen dia, senor.",
            "Bloqueado. No deje esperar demasiado.",
        ]
        return random.choice(options)

    def _handle_shutdown(self, match) -> str:
        result = pc_control.shutdown_pc("shutdown")
        prefixes = [
            "Como ordene, senor.",
            "Procedo con el apagado.",
            "Muy bien, senor.",
            "Entendido.",
        ]
        return f"{random.choice(prefixes)} {result['message']}"

    def _handle_restart(self, match) -> str:
        result = pc_control.shutdown_pc("restart")
        prefixes = [
            "Reiniciando. De nuevo.",
            "Muy bien, senor.",
            "A la orden.",
            "Procedo con el reinicio.",
        ]
        return f"{random.choice(prefixes)} {result['message']}"

    def _handle_sleep(self, match) -> str:
        result = pc_control.shutdown_pc("sleep")
        return f"{self._random_response('system')} {result['message']}"

    def _handle_hibernate(self, match) -> str:
        result = pc_control.shutdown_pc("hibernate")
        return f"{self._random_response('system')} {result['message']}"

    def _handle_logoff(self, match) -> str:
        result = pc_control.shutdown_pc("logoff")
        return f"{self._random_response('system')} {result['message']}"

    def _handle_cancel_shutdown(self, match) -> str:
        result = pc_control.shutdown_pc("cancel")
        return f"{self._random_response('system')} {result['message']}"

    def _handle_type_text(self, match) -> str | None:
        text = match.group(1).strip()
        if not text:
            return None
        try:
            import pyautogui
            import time
            time.sleep(0.5)
            pyautogui.typewrite(text, interval=0.02) if text.isascii() else pyautogui.write(text)
            return f"{self._random_response('system')} Texto escrito."
        except ImportError:
            return None  # Dejar a la IA
        except Exception:
            return None

    def _handle_click_info(self, match) -> str:
        return "No puedo identificar donde hacer click sin ver la pantalla. Puedo hacer click en coordenadas especificas si me das la posicion X,Y."

    def _handle_press_key(self, match) -> str | None:
        key = match.group(1).strip().lower()
        key_map = {
            "enter": "enter", "intro": "enter",
            "escape": "escape", "esc": "escape",
            "tab": "tab", "tabulador": "tab",
            "espacio": "space", "space": "space",
            "borrar": "backspace", "backspace": "backspace", "delete": "delete",
            "f1": "f1", "f2": "f2", "f3": "f3", "f4": "f4", "f5": "f5",
            "f11": "f11", "f12": "f12",
        }
        mapped = key_map.get(key)
        if mapped:
            try:
                import pyautogui
                pyautogui.press(mapped)
                return f"{self._random_response('system')} Tecla {key} presionada."
            except ImportError:
                return None
        return None

    # =========================================================================
    # Creacion de documentos y archivos
    # =========================================================================

    def _handle_create_word(self, match) -> str:
        content = match.group(1).strip().rstrip(".") if match.group(1) else ""
        desktop = os.path.expanduser("~\\Desktop").replace("\\", "\\\\")

        if content:
            # Limpiar el contenido para usarlo como titulo y texto
            safe_content = content.replace('"', '\\"').replace("'", "\\'")
            filename = re.sub(r'[^\w\s-]', '', content[:40]).strip().replace(' ', '_') or "documento"
        else:
            safe_content = "Documento creado por Jarvis"
            filename = "documento"

        code = f'''import subprocess
subprocess.run(["pip", "install", "python-docx"], capture_output=True)
from docx import Document
doc = Document()
doc.add_heading("{safe_content[:80]}", level=1)
doc.add_paragraph("{safe_content}")
filepath = r"{desktop}\\{filename}.docx"
doc.save(filepath)
import os
os.startfile(filepath)
print(f"Documento creado: {{filepath}}")
'''
        result = execute_code(code)
        if result["success"]:
            return f"Listo, senor. Documento Word creado en su escritorio: {filename}.docx"
        return f"No pude crear el documento: {result['message']}"

    def _handle_create_excel(self, match) -> str:
        content = match.group(1).strip().rstrip(".") if match.group(1) else ""
        desktop = os.path.expanduser("~\\Desktop").replace("\\", "\\\\")

        if content:
            safe_content = content.replace('"', '\\"')
            filename = re.sub(r'[^\w\s-]', '', content[:40]).strip().replace(' ', '_') or "hoja"
        else:
            safe_content = "Hoja de calculo"
            filename = "hoja"

        code = f'''import subprocess
subprocess.run(["pip", "install", "openpyxl"], capture_output=True)
from openpyxl import Workbook
wb = Workbook()
ws = wb.active
ws.title = "{safe_content[:30]}"
ws["A1"] = "{safe_content}"
filepath = r"{desktop}\\{filename}.xlsx"
wb.save(filepath)
import os
os.startfile(filepath)
print(f"Excel creado: {{filepath}}")
'''
        result = execute_code(code)
        if result["success"]:
            return f"Listo, senor. Excel creado en su escritorio: {filename}.xlsx"
        return f"No pude crear el Excel: {result['message']}"

    def _handle_create_txt(self, match) -> str:
        content = match.group(1).strip().rstrip(".") if match.group(1) else ""
        desktop = os.path.expanduser("~\\Desktop").replace("\\", "\\\\")

        if content:
            safe_content = content.replace('"', '\\"')
            filename = re.sub(r'[^\w\s-]', '', content[:40]).strip().replace(' ', '_') or "nota"
        else:
            safe_content = "Nota creada por Jarvis"
            filename = "nota"

        code = f'''filepath = r"{desktop}\\{filename}.txt"
with open(filepath, "w", encoding="utf-8") as f:
    f.write("{safe_content}")
import os
os.startfile(filepath)
print(f"Archivo creado: {{filepath}}")
'''
        result = execute_code(code)
        if result["success"]:
            return f"Listo, senor. Archivo de texto creado en su escritorio: {filename}.txt"
        return f"No pude crear el archivo: {result['message']}"

    # =========================================================================
    # Info del sistema y archivos
    # =========================================================================

    def _handle_list_files(self, match) -> str:
        folder = match.group(1).strip().lower().rstrip(".")
        folder_map = {
            "escritorio": "Desktop",
            "desktop": "Desktop",
            "descargas": "Downloads",
            "downloads": "Downloads",
            "documentos": "Documents",
            "documents": "Documents",
        }
        folder_name = folder_map.get(folder, folder)

        # Si es una ruta conocida, usar expanduser
        if folder_name in ("Desktop", "Downloads", "Documents"):
            path = os.path.expanduser(f"~\\{folder_name}")
        else:
            path = folder_name

        try:
            if not os.path.exists(path):
                return f"No encontre la carpeta '{folder}'."
            items = os.listdir(path)
            if not items:
                return f"La carpeta {folder} esta vacia."
            listing = "\n".join(f"  - {item}" for item in items[:30])
            total = len(items)
            extra = f"\n  ... y {total - 30} mas." if total > 30 else ""
            return f"Archivos en {folder} ({total}):\n{listing}{extra}"
        except Exception as e:
            return f"Error listando archivos: {e}"

    def _handle_disk_space(self, match) -> str:
        try:
            drive = match.group(1).strip().upper() if match.lastindex and match.group(1) else "C"
        except (IndexError, AttributeError):
            drive = "C"
        if not drive.endswith(":"):
            drive += ":"
        try:
            import shutil
            total, used, free = shutil.disk_usage(f"{drive}/")
            total_gb = total // (1024 ** 3)
            free_gb = free // (1024 ** 3)
            used_gb = used // (1024 ** 3)
            pct = round(used / total * 100, 1)
            return f"Disco {drive} tiene {free_gb} GB libres de {total_gb} GB totales. Usado: {used_gb} GB ({pct}%)."
        except Exception as e:
            return f"No pude obtener info del disco {drive}: {e}"

    def _handle_system_info(self, match) -> str:
        result = pc_control.get_system_info()
        if result["success"]:
            data = result["data"]
            return (
                f"Info del sistema:\n"
                f"  - OS: {data['os']}\n"
                f"  - Version: {data['version']}\n"
                f"  - Procesador: {data['processor']}\n"
                f"  - Usuario: {data['user']}\n"
                f"  - PC: {data['pc_name']}"
            )
        return "No pude obtener la info del sistema."

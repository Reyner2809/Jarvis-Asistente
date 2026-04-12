"""
Carga y chunking de documentos para el sistema RAG de Jarvis.

Soporta: PDF, DOCX, TXT, CSV, Markdown, y archivos de codigo.
Todas las dependencias de parseo (pypdf, python-docx) son opcionales:
si no estan instaladas, Jarvis informa al usuario y sigue funcionando.
"""

import csv
import io
import logging
import os
from pathlib import Path

log = logging.getLogger("jarvis.knowledge.loader")

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc",
    ".txt", ".md", ".csv",
    ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".cs",
    ".html", ".css", ".json", ".xml", ".yaml", ".yml",
    ".sh", ".bat", ".ps1", ".sql",
    ".go", ".rs", ".rb", ".php", ".swift", ".kt",
}


def is_supported(file_path: str) -> bool:
    return Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_text(file_path: str) -> str:
    """Extrae texto plano de un archivo segun su extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext == ".csv":
        return _extract_csv(file_path)
    else:
        return _extract_plain(file_path)


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Divide texto en fragmentos de ~chunk_size tokens (aprox palabras) con
    overlap entre fragmentos para mantener contexto en las fronteras.
    """
    if not text or not text.strip():
        return []

    words = text.split()
    if len(words) <= chunk_size:
        return [text.strip()]

    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap

    return chunks


def load_and_chunk(file_path: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Pipeline completo: extraer texto + chunquear."""
    text = extract_text(file_path)
    return chunk_text(text, chunk_size, overlap)


# ---------------------------------------------------------------------------
# Extractores por formato
# ---------------------------------------------------------------------------

def _extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "Para procesar PDFs instala pypdf: pip install pypdf"
        )
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "Para procesar DOCX instala python-docx: pip install python-docx"
        )
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tambien extraer tablas
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_csv(path: str) -> str:
    lines = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            lines.append(" | ".join(row))
    return "\n".join(lines)


def _extract_plain(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()
